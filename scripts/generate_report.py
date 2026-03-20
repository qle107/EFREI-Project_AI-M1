"""
Generate the final CineMatch project report as a professional .docx Word document.
Blue-themed, AI-focused, with Gantt chart, GitHub workflow, critique section.
"""
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT = Path(__file__).resolve().parent.parent / "docs" / "RAPPORT_FINAL_PROJET_IA_GENERATIVE.docx"

BLUE = RGBColor(0x1B, 0x4F, 0x72)
BLUE_LIGHT = RGBColor(0x21, 0x6F, 0xA3)
BLUE_ACCENT = RGBColor(0x2E, 0x86, 0xC1)
DARK = RGBColor(0x1C, 0x1C, 0x2E)
GRAY = RGBColor(0x5D, 0x6D, 0x7E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Global styles ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = DARK
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for lvl, sz in [(1, Pt(22)), (2, Pt(16)), (3, Pt(13))]:
    hs = doc.styles[f"Heading {lvl}"]
    hs.font.name = "Calibri"
    hs.font.size = sz
    hs.font.color.rgb = BLUE
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(18 if lvl == 1 else 12)
    hs.paragraph_format.space_after = Pt(6)

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def bold_run(para, text, *, color=None):
    r = para.add_run(text)
    r.font.bold = True
    if color:
        r.font.color.rgb = color
    return r


def add_para(text, *, bold_phrases=None, after_pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after_pt)
    if bold_phrases:
        remaining = text
        for bp in bold_phrases:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                p.add_run(remaining[:idx])
            bold_run(p, bp, color=BLUE)
            remaining = remaining[idx + len(bp):]
        if remaining:
            p.add_run(remaining)
    else:
        p.add_run(text)
    return p


def add_code(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1C, 0x1C, 0x2E)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB"/>')
    p._element.get_or_add_pPr().append(shading)


def add_img_placeholder(caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"[ INSERT IMAGE: {caption} ]")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = BLUE_ACCENT
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB"/>')
    p._element.get_or_add_pPr().append(shading)


def set_cell_bg(cell, hex_color: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def add_table(headers, rows, *, header_bg="1B4F72", header_fg=WHITE):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = header_fg
        r.font.name = "Calibri"
        set_cell_bg(cell, header_bg)
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg = "FFFFFF" if ri % 2 == 0 else "EBF5FB"
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = "Calibri"
            set_cell_bg(cell, bg)
    doc.add_paragraph()


def blue_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 60)
    run.font.color.rgb = BLUE_ACCENT
    run.font.size = Pt(8)


# ═══════════════════════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

add_img_placeholder("Logo EFREI Paris — fichier: Efrei-logo-couleur.svg")
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CineMatch")
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Moteur de Recommandation Cinématographique\nSémantique et Génératif")
r.font.size = Pt(16)
r.font.color.rgb = BLUE_LIGHT

doc.add_paragraph()
blue_separator()
doc.add_paragraph()

cover_lines = [
    ("Mastère Data Engineering et IA — EFREI Paris", False),
    ("Année universitaire 2025–2026", False),
    ("", False),
    ("CHUEMBOU MBAH Adrien Duval  &  QUAND DAT Le", True),
    ("", False),
    ("Tutrice : Sarah Malaeb", False),
    ("", False),
    ("Projet certifiant RNCP40875", True),
    ("Expert en Ingénierie de Données — Bloc 2 :", True),
    ("Pilotage et implémentation de solutions IA", True),
]
for text, is_bold in cover_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    if is_bold:
        r.font.bold = True
        r.font.color.rgb = BLUE

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS placeholder
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("Table des matières", level=1)
toc_items = [
    "2. Résumé Exécutif",
    "3. Introduction et Contexte",
    "4. Analyse du Besoin Utilisateur",
    "5. Méthodologie de Travail et Gestion de Projet",
    "6. Référentiel de Données",
    "7. Pipeline IA et Architecture",
    "8. Implémentation Technique",
    "9. Interface Utilisateur / Prototype",
    "10. Résultats et Tests",
    "11. Critique du Projet — Forces et Faiblesses",
    "12. Limites et Pistes d'Amélioration",
    "13. Conclusion",
    "14. Annexes",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.color.rgb = BLUE_LIGHT

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  2. RÉSUMÉ EXÉCUTIF
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Résumé Exécutif", level=1)

add_para(
    "CineMatch est un moteur de recommandation de films combinant NLP sémantique "
    "(Sentence-BERT, similarité cosinus, scoring multicritère pondéré) et IA Générative "
    "(LLM via Ollama / Claude / Gemini). Le système transforme une requête utilisateur "
    "en langage naturel en recommandations Top-3 accompagnées d'une explication structurée "
    "et d'un profil cinéphile généré.",
    bold_phrases=["NLP sémantique", "IA Générative", "Top-3", "profil cinéphile"],
)

add_para(
    "L'architecture repose sur deux phases : un pipeline offline (construction du référentiel "
    "sémantique + encodage SBERT de 9 841 films sur 4 axes indépendants) et un pipeline "
    "online (profilage utilisateur, cosine similarity, agrégation pondérée, ranking, "
    "enrichissement et explication GenAI). L'ensemble est exposé via une API FastAPI "
    "sécurisée (JWT) et un frontend Next.js.",
    bold_phrases=["pipeline offline", "pipeline online", "9 841 films", "4 axes indépendants"],
)

doc.add_heading("Technologies clés", level=3)
add_table(
    ["Composant", "Technologie", "Rôle"],
    [
        ["Embeddings", "all-MiniLM-L6-v2 (SBERT)", "Encodage sémantique dense 384-dim"],
        ["Similarité", "scikit-learn cosine_similarity", "Comparaison user ↔ corpus par axe"],
        ["Scoring", "Formule pondérée CineMatch", "0.35 mood + 0.25 theme + 0.20 style + 0.15 desc + 0.05 recency"],
        ["GenAI", "Ollama / Claude / Gemini", "Enrichissement texte court + explication + profil cinéphile"],
        ["Cache LLM", "SHA-256 in-memory, TTL 24h", "Réduction coût/latence (req 1.5.3)"],
        ["API", "FastAPI + JWT HS-256", "REST, OpenAPI, authentification"],
        ["Frontend", "Next.js 16 + React + Tailwind", "Prototype interactif, radar charts SVG"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. INTRODUCTION ET CONTEXTE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Introduction et Contexte", level=1)

doc.add_heading("3.1 Problématique", level=2)
add_para(
    "Les systèmes de recommandation classiques (filtrage collaboratif, content-based sur métadonnées) "
    "peinent à capturer la sémantique profonde des préférences utilisateur. Un spectateur qui demande "
    "\"un thriller psychologique sombre avec des dilemmes moraux\" ne peut être satisfait par un simple "
    "filtre Genre = Thriller. Il faut comprendre le sens du texte, pas seulement les mots-clés.",
    bold_phrases=["sémantique profonde", "comprendre le sens du texte"],
)

add_para(
    "CineMatch résout ce problème en projetant les films et les requêtes dans un espace vectoriel "
    "sémantique partagé grâce à Sentence-BERT, puis en mesurant la proximité sur 4 axes "
    "indépendants : Mood (atmosphère), Theme (genre/thème), Style (rythme narratif), "
    "et Description (texte libre). Ce multi-axe permet un scoring fin, explicable par radar chart.",
    bold_phrases=["espace vectoriel sémantique partagé", "Sentence-BERT", "4 axes indépendants",
                  "Mood", "Theme", "Style", "Description"],
)

doc.add_heading("3.2 Pourquoi cette approche IA", level=2)
add_para(
    "Le choix d'une approche par embeddings sémantiques plutôt que par filtrage classique "
    "se justifie par plusieurs raisons fondamentales :",
)

rationale = [
    ("Représentation dense vs. sparse",
     "Les embeddings SBERT produisent des vecteurs denses de 384 dimensions qui capturent "
     "les relations sémantiques entre concepts (\"sombre\" ≈ \"noir\" ≈ \"dark\"), là où un "
     "bag-of-words ou TF-IDF ne verrait aucun lien."),
    ("Comparabilité universelle",
     "En encodant utilisateurs et films dans le même espace vectoriel, la similarité cosinus "
     "devient une métrique naturelle de pertinence, sans besoin de règles ad-hoc."),
    ("Explicabilité par axe",
     "Le découpage en 4 dimensions sémantiques indépendantes (plutôt qu'un unique vecteur "
     "combiné) permet d'expliquer pourquoi un film est recommandé : \"score mood élevé mais "
     "style moyen\". C'est la base du radar chart."),
    ("GenAI contrôlée (RAG-like)",
     "La couche générative n'invente pas : elle reçoit les Top-3 calculés mathématiquement "
     "et les scores par axe, puis produit une explication grounded. Cela limite les "
     "hallucinations et le coût."),
]
for title, desc in rationale:
    p = doc.add_paragraph()
    bold_run(p, f"{title} — ", color=BLUE)
    p.add_run(desc)

doc.add_heading("3.3 Alternatives considérées", level=2)
add_table(
    ["Approche", "Avantages", "Pourquoi non retenue"],
    [
        ["Filtrage collaboratif", "Pas besoin de features film", "Nécessite un historique utilisateur massif (cold start)"],
        ["TF-IDF + cosinus", "Simple, rapide", "Représentation sparse, pas de capture sémantique"],
        ["all-mpnet-base-v2", "Meilleur STS score (86.8)", "4× plus lent, 768 dims, overkill pour texte court cinéma"],
        ["Fine-tuned SBERT", "Qualité maximale", "Nécessite un dataset annoté (paires query/film), hors scope"],
        ["LLM-only (sans retrieval)", "Simplicité", "Coûteux, hallucinations non contrôlées, pas de scoring"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. ANALYSE DU BESOIN UTILISATEUR
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Analyse du Besoin Utilisateur", level=1)

doc.add_heading("4.1 Persona cible", level=2)
add_para(
    "L'utilisateur type est un cinéphile occasionnel à confirmé, cherchant une recommandation "
    "personnalisée sans avoir de titre précis en tête. Il exprime ses envies via une combinaison de "
    "texte libre (\"je veux un film sombre avec des rebondissements\") et de choix structurés "
    "(mood, genre, style, intensité via sliders Likert 1-5).",
    bold_phrases=["texte libre", "choix structurés", "sliders Likert 1-5"],
)

doc.add_heading("4.2 Scénarios d'usage", level=2)
add_table(
    ["Scénario", "Input", "Comportement IA"],
    [
        ["A — Détaillé", "\"Thriller SF sombre avec dilemmes moraux IA, tense 5/5\"",
         "Encoding direct SBERT → cosinus → Top-3 → explication"],
        ["B — Court (< 5 mots)", "\"film triste\"",
         "Enrichissement GenAI → nouvel encoding → cosinus → Top-3"],
        ["C — Preset démo", "Clic sur 'Dark Crime Thriller'",
         "Pipeline complet, cache LLM après 1er appel → réponse instantanée"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. MÉTHODOLOGIE ET GESTION DE PROJET
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Méthodologie de Travail et Gestion de Projet", level=1)

doc.add_heading("5.1 Approche adoptée", level=2)
add_para(
    "Nous avons suivi une approche hybride Kanban/Agile, avec des incréments fonctionnels "
    "courts et des validations régulières. Chaque lot produit un livrable testable "
    "(pipeline, API, UI, GenAI, historique, settings, documentation).",
    bold_phrases=["hybride Kanban/Agile", "incréments fonctionnels"],
)

doc.add_heading("5.2 Diagramme de Gantt — Planning du projet", level=2)
add_para(
    "Le tableau ci-dessous résume le planning par lot de travail, réparti sur ~8 semaines :"
)
add_table(
    ["Lot", "Tâche", "Sem 1-2", "Sem 3-4", "Sem 5-6", "Sem 7-8"],
    [
        ["1", "Nettoyage données + référentiel", "████", "", "", ""],
        ["2", "Pipeline embeddings SBERT", "████", "", "", ""],
        ["3", "Scoring multicritère pondéré", "", "████", "", ""],
        ["4", "API REST FastAPI + JWT", "", "████", "", ""],
        ["5", "Couche GenAI (prompts, cache, multi-provider)", "", "", "████", ""],
        ["6", "Frontend Next.js (questionnaire, résultats, radar)", "", "", "████", ""],
        ["7", "Presets, historique, settings runtime", "", "", "", "████"],
        ["8", "Documentation, rapport, soutenance", "", "", "", "████"],
    ],
)

add_img_placeholder("Diagramme de Gantt visuel (si version graphique disponible)")

doc.add_heading("5.3 Collaboration en binôme avec GitHub", level=2)
add_para(
    "Le travail a été organisé sur un dépôt GitHub partagé. La collaboration a suivi "
    "ces pratiques :",
    bold_phrases=["dépôt GitHub partagé"],
)

collab_points = [
    ("Versionning Git",
     "Tous les fichiers source (backend, frontend, scripts, data processing) sont versionnés. "
     "Chaque fonctionnalité majeure (pipeline, API, GenAI, settings) a fait l'objet de commits distincts."),
    ("Organisation modulaire",
     "Le code est structuré en modules Python séparés (src/preprocessing, src/embedding, "
     "src/scoring, src/genAi, src/api) pour éviter les conflits de merge."),
    ("Répartition du travail",
     "Le binôme a réparti les tâches par domaine : un membre focalisé sur le pipeline IA "
     "(embeddings, scoring, GenAI), l'autre sur l'API et le frontend."),
    ("Documentation continue",
     "Le fichier PROJECT_DOCUMENTATION.md a été maintenu à jour tout au long du projet, "
     "servant de référence partagée entre les deux membres."),
    (".env et configuration",
     "Les variables sensibles (clés API, secrets JWT) sont gérées via .env avec un "
     ".env.example versionné pour faciliter l'onboarding."),
]
for title, desc in collab_points:
    p = doc.add_paragraph()
    bold_run(p, f"• {title} — ", color=BLUE)
    p.add_run(desc)

doc.add_heading("5.4 Gestion des risques", level=2)
add_table(
    ["Risque identifié", "Impact", "Mitigation mise en place"],
    [
        ["Latence LLM (Ollama sur CPU)", "Élevé", "Cache SHA-256 (TTL 24h) + switch vers Claude/Gemini cloud"],
        ["Indisponibilité API cloud", "Moyen", "Fallback automatique vers Ollama local"],
        ["Bruit sémantique (labels 'neutral')", "Moyen", "Décomposition 4 axes + pondération différenciée"],
        ["Drift de qualité des embeddings", "Faible", "Monitoring logs, possibilité de re-build offline"],
        ["Conflits Git en binôme", "Faible", "Modules séparés, communication régulière"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. RÉFÉRENTIEL DE DONNÉES
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Référentiel de Données", level=1)

doc.add_heading("6.1 Source et volume", level=2)
add_para(
    "Le corpus est un dataset dérivé de TMDB contenant 9 841 films avec les champs : "
    "Title, Overview (synopsis), Genre, Release_Date, Vote_Average, Poster_Url. "
    "Ce dataset a été nettoyé (suppression doublons, extraction année, normalisation genres) "
    "pour produire Database_Cleaned.csv.",
    bold_phrases=["9 841 films", "Database_Cleaned.csv"],
)

doc.add_heading("6.2 Extraction des blocs sémantiques", level=2)
add_para(
    "Le module block_builder.py transforme chaque overview de film en 4 blocs sémantiques "
    "via une taxonomie de mots-clés avec matching regex (word-boundary). Ce processus est "
    "déterministe et reproductible.",
    bold_phrases=["4 blocs sémantiques", "taxonomie de mots-clés"],
)

add_table(
    ["Bloc", "Source", "Taxonomie", "Exemples de labels"],
    [
        ["Mood", "Overview", "MOOD_KEYWORDS", "dark, uplifting, tense, heroic"],
        ["Theme", "Overview", "THEME_KEYWORDS", "crime, fantasy, survival, family"],
        ["NarrativeStyle", "Overview", "STYLE_KEYWORDS", "action, mystery, drama"],
        ["Description", "Overview + Genre", "—", "Texte complet enrichi du genre"],
    ],
)

doc.add_heading("Comment fonctionne l'extraction", level=3)
add_para(
    "Pour chaque film, l'overview est converti en minuscules puis scanné avec des regex "
    "\\b{keyword}\\b (word boundary). Si un mot-clé d'une catégorie est détecté, le label "
    "correspondant est ajouté. Un film peut avoir plusieurs labels par bloc (ex: \"dark, tense\"). "
    "Si aucun mot-clé ne matche, le label par défaut est \"neutral\".",
    bold_phrases=["regex", "word boundary", "\"neutral\""],
)

add_code("""def extract_block(text, taxonomy):
    text = clean_text(text)  # lowercase + handle NaN
    detected = []
    for label, keywords in taxonomy.items():
        for word in keywords:
            if re.search(rf"\\b{word}\\b", text):
                detected.append(label)
                break  # one match per label is enough
    return ", ".join(detected) if detected else "neutral"

# Application aux 9 841 films :
df["Mood"]  = df["Overview"].apply(lambda x: extract_block(x, MOOD_KEYWORDS))
df["Theme"] = df["Overview"].apply(lambda x: extract_block(x, THEME_KEYWORDS))
df["NarrativeStyle"] = df["Overview"].apply(
    lambda x: extract_block(x, STYLE_KEYWORDS)
)
df["Description"] = df["Overview"] + " Genre: " + df["Genre"]""")

doc.add_heading("6.3 Schéma du référentiel produit", level=2)
add_table(
    ["Colonne", "Type", "Description"],
    [
        ["FilmID", "int", "Identifiant unique (index 0-based)"],
        ["Title", "str", "Titre du film"],
        ["Mood", "str", "Labels mood (ex: \"dark, tense\")"],
        ["Theme", "str", "Labels thématiques (ex: \"crime, family\")"],
        ["NarrativeStyle", "str", "Labels style (ex: \"mystery, action\")"],
        ["Description", "str", "Overview + \" Genre: \" + Genre"],
        ["release_year", "int", "Année de sortie"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. PIPELINE IA ET ARCHITECTURE (SECTION MAJEURE)
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Pipeline IA et Architecture", level=1)

add_para(
    "Le cœur technique de CineMatch repose sur un pipeline en 7 étapes, séparé en deux phases : "
    "une phase offline (exécutée une seule fois pour construire l'index) et une phase online "
    "(exécutée à chaque requête utilisateur). Chaque étape est détaillée ci-dessous avec "
    "le code correspondant, les choix techniques et leurs justifications.",
    bold_phrases=["pipeline en 7 étapes", "phase offline", "phase online"],
)

add_img_placeholder(
    "Diagramme pipeline complet : Questionnaire → Preprocessing → "
    "SBERT Encoding → Cosine Similarity → Weighted Scoring → Top-3 → GenAI Explanation"
)

# ── 7.1 OFFLINE ──
doc.add_heading("7.1 Phase Offline — Construction de l'index sémantique", level=2)

doc.add_heading("Étape 1 : Extraction des blocs sémantiques", level=3)
add_para(
    "Voir section 6.2. Le referential_builder.py lit Database_Cleaned.csv, applique "
    "block_builder.build_semantic_blocks(), et sauvegarde movies_referential.csv.",
)

doc.add_heading("Étape 2 : Encodage SBERT du corpus (Embedding Builder)", level=3)
add_para(
    "Chaque bloc de chaque film est encodé indépendamment avec le modèle SBERT "
    "all-MiniLM-L6-v2 pour produire un vecteur dense de 384 dimensions. "
    "Le résultat est 4 matrices NumPy de forme (9 841 × 384), une par axe sémantique.",
    bold_phrases=["all-MiniLM-L6-v2", "384 dimensions", "4 matrices NumPy", "(9 841 × 384)"],
)

add_para(
    "Pourquoi encoder chaque axe séparément ? Un unique embedding par film "
    "(ex: concaténation title+overview+genre) perdrait la granularité nécessaire au "
    "scoring par axe et à la visualisation radar. En séparant les axes, on peut "
    "pondérer différemment mood (35%) vs. description (15%) et expliquer "
    "précisément où le film matche ou diverge.",
    bold_phrases=["Pourquoi encoder chaque axe séparément ?", "pondérer différemment"],
)

add_code("""# src/embedding/embedding_builder.py
from sentence_transformers import SentenceTransformer
import numpy as np

model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

# Encode each semantic axis independently
mood_emb  = model.encode(df["Mood"].tolist(),  convert_to_numpy=True)
theme_emb = model.encode(df["Theme"].tolist(), convert_to_numpy=True)
style_emb = model.encode(df["NarrativeStyle"].tolist(), convert_to_numpy=True)
desc_emb  = model.encode(df["Description"].tolist(),    convert_to_numpy=True)

# Save as NumPy arrays for fast loading at inference
np.save("models/embeddings/mood_embeddings.npy",  mood_emb)   # (9841, 384)
np.save("models/embeddings/theme_embeddings.npy", theme_emb)  # (9841, 384)
np.save("models/embeddings/style_embeddings.npy", style_emb)  # (9841, 384)
np.save("models/embeddings/desc_embeddings.npy",  desc_emb)   # (9841, 384)""")

add_para(
    "Le modèle SBERT fonctionne ainsi : le texte est tokenisé, passé dans un backbone "
    "MiniLM (6 couches transformer), puis mean-pooled pour produire un vecteur de taille fixe. "
    "Des textes sémantiquement proches sont projetés vers des régions voisines de l'espace 384-dim.",
    bold_phrases=["MiniLM", "mean-pooled", "vecteur de taille fixe"],
)

# ── 7.2 ONLINE ──
doc.add_heading("7.2 Phase Online — Recommandation temps-réel", level=2)

doc.add_heading("Étape 3 : Encodage du profil utilisateur", level=3)
add_para(
    "Le questionnaire hybride de l'utilisateur (texte libre + choix structurés + Likert) "
    "est traduit en 4 blocs textuels selon des templates structurés. Chaque bloc est ensuite "
    "encodé avec le même modèle SBERT que le corpus, garantissant la comparabilité dans "
    "le même espace vectoriel.",
    bold_phrases=["4 blocs textuels", "même modèle SBERT", "même espace vectoriel"],
)

add_code("""# src/user_profile/profile_encoder.py
def build_profile_blocks(self, questionnaire):
    mood_block = (
        f"The user wants a {questionnaire.preferred_mood} mood "
        f"with intensity level {questionnaire.mood_intensity}"
    )
    theme_block = (
        f"The user prefers {questionnaire.preferred_genre} themes "
        f"with interest level {questionnaire.theme_interest}"
    )
    style_block = (
        f"The user enjoys {questionnaire.preferred_style} narrative "
        f"pacing with interest level {questionnaire.style_interest}"
    )
    desc_block = questionnaire.description
    if questionnaire.preferred_era:
        desc_block += f" Preferred era: {questionnaire.preferred_era}."
    if questionnaire.preferred_director:
        desc_block += f" Style of: {questionnaire.preferred_director}."
    return {"mood": mood_block, "theme": theme_block,
            "style": style_block, "description": desc_block}

# Encode with same SBERT model
profile = {k: model.encode(v) for k, v in blocks.items()}""")

add_para(
    "L'intensité Likert (ex: \"intensity level 5\" vs \"intensity level 1\") influence "
    "la position du vecteur dans l'espace sémantique. Bien que subtile, cette variation "
    "modifie les scores de similarité de manière cohérente avec l'intention utilisateur.",
    bold_phrases=["intensité Likert", "influence la position du vecteur"],
)

doc.add_heading("Étape 4 : Calcul de similarité cosinus", level=3)
add_para(
    "Pour chaque axe, on calcule la similarité cosinus entre le vecteur utilisateur "
    "et les 9 841 vecteurs films correspondants. La formule est :",
    bold_phrases=["similarité cosinus"],
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("sim(A, B) = (A · B) / (‖A‖ × ‖B‖)")
r.font.bold = True
r.font.size = Pt(12)
r.font.color.rgb = BLUE

add_para(
    "Le résultat est borné dans [−1, 1] (typiquement [0, 1] pour des embeddings normalisés). "
    "Plus la valeur est élevée, plus le film est sémantiquement proche du profil sur cet axe.",
    bold_phrases=["[−1, 1]", "sémantiquement proche"],
)

add_code("""# src/scoring/coverage_scorer.py
from sklearn.metrics.pairwise import cosine_similarity

# 4 comparaisons indépendantes : user vector vs corpus matrix
mood_sim  = cosine_similarity([user_profile["mood"]],  self.mood_emb)[0]
theme_sim = cosine_similarity([user_profile["theme"]], self.theme_emb)[0]
style_sim = cosine_similarity([user_profile["style"]], self.style_emb)[0]
desc_sim  = cosine_similarity([user_profile["description"]], self.desc_emb)[0]
# Chaque résultat : array de 9 841 scores""")

doc.add_heading("Étape 5 : Normalisation min-max par axe", level=3)
add_para(
    "Avant agrégation, chaque axe est normalisé indépendamment dans [0, 1] "
    "via min-max scaling. Cela garantit que les axes avec des distributions "
    "naturellement différentes contribuent équitablement au score final.",
    bold_phrases=["normalisé indépendamment", "[0, 1]", "min-max scaling"],
)

add_code("""def normalize(self, sim):
    return (sim - sim.min()) / (sim.max() - sim.min() + 1e-8)
    # epsilon 1e-8 pour éviter division par zéro""")

doc.add_heading("Étape 6 : Agrégation pondérée — Formule CineMatch", level=3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
r = p.add_run(
    "CoverageScore = 0.35 × Mood + 0.25 × Theme + 0.20 × Style + 0.15 × Desc + 0.05 × Recency"
)
r.font.bold = True
r.font.size = Pt(12)
r.font.color.rgb = BLUE
shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB"/>')
p._element.get_or_add_pPr().append(shading)

add_code("""# Weighted aggregation
final_score = (
    0.35 * mood_sim +     # Atmosphere is most discriminating
    0.25 * theme_sim +    # Genre/theme is second filter
    0.20 * style_sim +    # Narrative pacing refines matching
    0.15 * desc_sim +     # Free-text captures nuances
    0.05 * recency_score  # Light tiebreaker for freshness
)""")

doc.add_heading("Justification détaillée des poids", level=3)
add_table(
    ["Axe", "Poids", "Pourquoi"],
    [
        ["Mood", "35%", "L'atmosphère émotionnelle est le critère le plus discriminant : un film \"dark\" vs \"uplifting\" change radicalement l'expérience."],
        ["Theme", "25%", "Le genre est le deuxième filtre naturel : un amateur de SF ne veut pas de romance."],
        ["Style", "20%", "Le rythme narratif affine : un thriller \"slow-burn\" vs \"fast-paced\" n'est pas interchangeable."],
        ["Description", "15%", "Le texte libre capture des nuances non couvertes par les axes structurés (directeur, ère, ton spécifique)."],
        ["Recency", "5%", "Tie-breaker léger : exp(−Δyear/10). Un film de 10 ans score ~0.37, 30 ans ~0.05. Ne domine jamais."],
    ],
)

add_para(
    "Le recency score est calculé par décroissance exponentielle :",
    bold_phrases=["décroissance exponentielle"],
)
add_code("""recency_score = np.exp(-(current_year - release_year) / 10)""")

doc.add_heading("Étape 7 : Ranking Top-3 et output", level=3)
add_code("""ranked = self.df.sort_values(by="CoverageScore", ascending=False)
top3 = ranked.head(3)
# Each movie includes: FilmID, Title, CoverageScore,
# MoodScore, ThemeScore, StyleScore, DescScore""")

add_para(
    "Les 3 films avec le CoverageScore le plus élevé sont retournés avec leurs sous-scores "
    "détaillés. Ces sous-scores alimentent le radar chart SVG dans le frontend et permettent "
    "à l'utilisateur de comprendre sur quel axe le film matche le mieux.",
    bold_phrases=["sous-scores détaillés", "radar chart SVG"],
)

doc.add_page_break()

# ── 7.3 GENAI ──
doc.add_heading("7.3 Couche GenAI — Enrichissement et Explication", level=2)

add_para(
    "La couche générative intervient de manière contrôlée à deux moments précis du pipeline. "
    "Elle ne remplace jamais le scoring mathématique : elle le complète en langage naturel.",
    bold_phrases=["contrôlée", "ne remplace jamais le scoring mathématique"],
)

doc.add_heading("Enrichissement de texte court (< 5 mots)", level=3)
add_para(
    "Quand l'utilisateur saisit une description très courte (ex: \"film triste\"), "
    "un unique appel LLM l'enrichit en 2-3 phrases avant l'encodage SBERT. "
    "Cela améliore la qualité du vecteur description sans modifier les autres axes.",
    bold_phrases=["unique appel LLM", "avant l'encodage SBERT"],
)

add_code("""# src/genAi/prompt_builder.py
def build_enrichment_prompt(short_description: str) -> str:
    return (
        "You are a movie preference assistant. The user typed a very "
        "brief description. Expand this into a detailed 2-3 sentence "
        "description that preserves the original intent and adds "
        "useful context about tone, themes, pacing, and style.\\n\\n"
        f'User input: "{short_description}"\\n\\n'
        "Expanded description (2-3 sentences, no preamble):"
    )""")

doc.add_heading("Explication structurée + Profil cinéphile", level=3)
add_para(
    "Après le ranking Top-3, un unique appel LLM reçoit les préférences utilisateur "
    "et les 3 films avec leurs scores par axe. Le prompt demande une sortie structurée "
    "en sections : OVERVIEW, FILM_1, FILM_2, FILM_3, REFINE, CINEPHILE_PROFILE.",
    bold_phrases=["unique appel LLM", "sortie structurée"],
)

add_code("""# Prompt structure (condensed)
prompt = (
    "You are a cinematic AI analyst. Write a structured explanation "
    "using EXACTLY these headers:\\n"
    "OVERVIEW: [why these 3 match]\\n"
    "FILM_1: {title1} [strongest dimensions]\\n"
    "FILM_2: {title2} [...]\\n"
    "FILM_3: {title3} [...]\\n"
    "REFINE: [how to get better results]\\n"
    "CINEPHILE_PROFILE: [2-sentence taste profile]\\n\\n"
    "USER PREFERENCES: ...\\n"
    "SEMANTIC SCORES: ...\\n"
    "RULES: Do NOT invent data. Stay grounded in scores."
)""")

doc.add_heading("Cache LLM — Réduction coût et latence", level=3)
add_para(
    "Chaque prompt est hashé via SHA-256 et la réponse stockée en mémoire avec TTL (24h) "
    "et taille max (200 entrées). Quand un prompt identique est resoumis, la réponse est "
    "servie instantanément depuis le cache. Ce mécanisme est essentiel pour les presets "
    "et les démonstrations répétées.",
    bold_phrases=["SHA-256", "TTL (24h)", "200 entrées", "servie instantanément"],
)

add_code("""# src/genAi/llm_client.py — PromptCache (simplified)
class _PromptCache:
    def __init__(self, ttl=86400, max_size=200):
        self._store = {}  # {sha256_key: (response, timestamp)}

    def get(self, prompt):
        key = sha256(prompt)
        if key in store and not expired:
            return cached_response  # CACHE HIT

    def put(self, prompt, response):
        if at_capacity:
            evict_oldest()
        store[key] = (response, now())""")

doc.add_heading("Multi-provider LLM", level=3)
add_table(
    ["Provider", "Modèle", "Latence typique", "Coût", "Usage"],
    [
        ["Ollama (local)", "phi3:mini", "2-10s (CPU)", "Gratuit", "Développement, pas d'internet"],
        ["Claude (Anthropic)", "claude-3-5-haiku", "1-3s", "API payante", "Qualité supérieure"],
        ["Gemini (Google)", "gemini-1.5-flash", "1-2s", "API payante", "Rapide, bonne qualité"],
    ],
)

add_para(
    "Le provider actif est switchable à chaud depuis la page Settings ou le sélecteur "
    "\"Explanations by\" sur la page Recommend, sans redémarrer le serveur.",
    bold_phrases=["switchable à chaud", "sans redémarrer"],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. IMPLÉMENTATION TECHNIQUE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Implémentation Technique", level=1)

doc.add_heading("8.1 Comparaison des modèles SBERT", level=2)
add_table(
    ["Critère", "all-MiniLM-L6-v2 ✓", "all-mpnet-base-v2", "multi-qa-MiniLM"],
    [
        ["Paramètres", "22M", "109M", "22M"],
        ["Dimension vecteur", "384", "768", "384"],
        ["STS Benchmark", "82.6", "86.8", "81.0"],
        ["Vitesse CPU", "~4000 sent/s", "~1000 sent/s", "~4000 sent/s"],
        ["Usage optimal", "Texte court/moyen", "Haute qualité STS", "Q&A"],
        ["Adapté au projet ?", "Oui (overviews 50-200 mots)", "Trop lent pour 9841 films", "Non (orienté Q&A)"],
    ],
)

doc.add_heading("8.2 Questionnaire hybride → vecteurs IA", level=2)
add_table(
    ["Champ utilisateur", "Type", "Transformation IA"],
    [
        ["description", "Texte libre (min 10 car.)", "Encodage SBERT direct → axe Description"],
        ["preferred_mood", "Chip sélection", "Template \"The user wants a {mood} mood...\" → axe Mood"],
        ["preferred_genre", "Chip sélection", "Template \"The user prefers {genre} themes...\" → axe Theme"],
        ["preferred_style", "Chip sélection", "Template \"The user enjoys {style} pacing...\" → axe Style"],
        ["mood_intensity", "Likert 1-5", "Intégré dans le template mood → influence le vecteur"],
        ["theme_interest", "Likert 1-5", "Intégré dans le template theme → influence le vecteur"],
        ["style_interest", "Likert 1-5", "Intégré dans le template style → influence le vecteur"],
        ["preferred_era", "Optionnel", "Concaténé au desc_block → influence axe Description"],
        ["preferred_director", "Optionnel", "Concaténé au desc_block → influence axe Description"],
    ],
)

doc.add_heading("8.3 Endpoints API principaux", level=2)
add_table(
    ["Méthode", "Endpoint", "Auth", "Description"],
    [
        ["POST", "/api/v1/recommendations", "JWT", "Top-3 + explication GenAI"],
        ["GET", "/api/v1/recommendations/presets", "—", "Liste presets disponibles"],
        ["POST", "/api/v1/recommendations/presets/{id}", "JWT", "Exécuter un preset"],
        ["GET", "/api/v1/recommendations/history", "JWT", "Historique par utilisateur"],
        ["GET/PUT", "/api/v1/settings/llm", "JWT", "Lire/modifier config LLM"],
        ["POST", "/api/v1/settings/llm/clear-cache", "JWT", "Purger cache LLM"],
    ],
)

doc.add_heading("8.4 Gouvernance et usage responsable", level=2)
gov_points = [
    "Cache LLM avec TTL et taille max — contrôle coût et latence",
    "Authentification JWT sur tous les endpoints sensibles",
    "Historisation complète des générations dans SQLite — traçabilité",
    "Logs structurés des appels LLM (provider, latence, cache hit/miss)",
    "Runtime switch provider sans redémarrage — résilience",
    "Séparation nette données / logique IA / service API / interface",
    "Clés API jamais retournées par l'API (stockées en mémoire uniquement)",
]
for point in gov_points:
    p = doc.add_paragraph()
    bold_run(p, "• ", color=BLUE)
    p.add_run(point)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. INTERFACE UTILISATEUR
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Interface Utilisateur — Prototype", level=1)

add_para(
    "Le prototype CineMatch est construit avec Next.js 16, React et Tailwind CSS. "
    "L'interface guide l'utilisateur à travers un questionnaire en 5 étapes, affiche "
    "les résultats Top-3 avec radar charts, explication GenAI et profil cinéphile.",
    bold_phrases=["Next.js 16", "5 étapes", "radar charts"],
)

screens = [
    "Page d'accueil CineMatch — stats (9841 films, 384-dim, 4 axes) + login",
    "Questionnaire de recommandation — description libre + mood/genre/style chips + sliders Likert",
    "Sélection provider LLM (Ollama/Claude/Gemini) en haut de la page Recommend",
    "Résultats Top-3 — affiches films, CoverageScore, explication structurée, profil cinéphile",
    "Radar charts SVG par film — Mood, Theme, Style, Description",
    "Page How It Works — architecture complète du pipeline, code extraits, scoring formula",
    "Page Settings — choix provider, température, modèle, cache TTL/size, purge cache",
    "Historique des recommandations — liste des générations passées avec timestamp",
]
for s in screens:
    add_img_placeholder(s)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  10. RÉSULTATS ET TESTS
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Résultats et Tests", level=1)

doc.add_heading("10.1 Démonstrations par profil", level=2)
add_table(
    ["Profil test", "Mood", "Genre", "Intensité", "Top-1 type attendu", "Enrichi ?", "Caché ?"],
    [
        ["Dark Crime Thriller", "dark", "crime", "5/5", "Film noir/thriller", "Non", "Oui (2e appel)"],
        ["Feel-Good Family", "uplifting", "family", "3/5", "Film familial/aventure", "Non", "Oui (2e appel)"],
        ["Mind-Bending Sci-Fi", "tense", "sci-fi", "4/5", "SF cérébral", "Non", "Non"],
        ["\"sad movie\" (court)", "—", "—", "—", "Drame émotionnel", "Oui", "Non"],
    ],
)

doc.add_heading("10.2 Métadonnées retournées par l'API", level=2)
indicators = [
    ("coverage_score / mood_score / theme_score / style_score / desc_score",
     "Scores numériques [0, 1] par axe — alimentent le radar chart"),
    ("explanation", "Texte GenAI structuré (OVERVIEW, FILM_1/2/3, REFINE) grounded dans les scores"),
    ("cinephile_profile", "2 phrases décrivant le goût de l'utilisateur"),
    ("description_enriched", "true si le texte court a été enrichi par GenAI avant encoding"),
    ("cached", "true si l'explication provient du cache (pas de nouvel appel LLM)"),
    ("llm_provider", "Provider utilisé : ollama / anthropic / gemini"),
]
for name, desc in indicators:
    p = doc.add_paragraph()
    bold_run(p, f"{name} — ", color=BLUE)
    p.add_run(desc)

add_img_placeholder("Radar charts comparés pour 3 profils différents (dark vs family vs sci-fi)")
add_img_placeholder("Exemple de réponse JSON avec tous les champs (coverage_score, explanation, etc.)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  11. CRITIQUE DU PROJET (NOUVELLE SECTION)
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Critique du Projet — Forces et Faiblesses", level=1)

doc.add_heading("11.1 Points forts", level=2)
strengths = [
    ("Pipeline IA complet bout-en-bout",
     "De la donnée brute CSV à la recommandation expliquée par GenAI, le projet couvre "
     "l'intégralité de la chaîne : preprocessing → embeddings → scoring → ranking → "
     "explication. C'est un atout majeur pour démontrer la maîtrise du pipeline ML."),
    ("Scoring explicable et multi-axe",
     "Le découpage en 4 axes indépendants (Mood, Theme, Style, Description) rend le score "
     "interprétable. Le radar chart permet à l'utilisateur de comprendre où le match est fort "
     "et où il est faible. Peu de systèmes de recommandation offrent ce niveau de transparence."),
    ("GenAI grounded et contrôlée",
     "L'IA générative ne fonctionne jamais en roue libre : elle reçoit les scores calculés "
     "et les films sélectionnés mathématiquement. Le prompt inclut des RULES explicites "
     "(\"Do NOT invent data\"). C'est une implémentation RAG responsable."),
    ("Cache LLM bien pensé",
     "Le mécanisme SHA-256 + TTL + éviction résout un problème réel de coût et latence. "
     "Les presets bénéficient particulièrement de ce cache pour des démos fluides."),
    ("Multi-provider switchable",
     "Pouvoir basculer entre Ollama (local/gratuit), Claude et Gemini en temps réel, "
     "sans redémarrage, est un avantage opérationnel significatif."),
    ("Interface utilisateur soignée",
     "Le questionnaire en 5 étapes, les radar charts SVG, le profil cinéphile, "
     "la page How It Works et les Settings forment un prototype complet."),
]
for title, desc in strengths:
    p = doc.add_paragraph()
    bold_run(p, f"✓ {title} — ", color=BLUE)
    p.add_run(desc)

doc.add_heading("11.2 Points faibles et axes d'amélioration", level=2)
weaknesses = [
    ("Taxonomie lexicale limitée",
     "L'extraction de blocs repose sur des regex mots-clés. Un overview sans keyword "
     "connu reçoit le label \"neutral\", ce qui réduit la discrimination. "
     "Avec plus de temps : NER ou classification fine-tuned remplacerait cette heuristique."),
    ("Pas de fine-tuning SBERT",
     "Le modèle est utilisé out-of-the-box. Un fine-tuning sur des paires "
     "(requête utilisateur, film pertinent) améliorerait significativement la qualité. "
     "Avec plus de ressources : constitution d'un dataset annoté + training."),
    ("Évaluation majoritairement qualitative",
     "Nous n'avons pas implémenté de métriques supervisées (Precision@K, NDCG, MRR). "
     "L'évaluation repose sur des tests manuels et l'inspection des résultats. "
     "Avec plus de temps : benchmark formel avec annotateurs."),
    ("Dataset statique et anglophone",
     "Les 9 841 films sont figés et en anglais. Pas d'ingestion temps-réel ni de support multilingue. "
     "Avec plus de ressources : API TMDB live + modèle SBERT multilingue."),
    ("Pas de base vectorielle dédiée",
     "Les embeddings sont stockés en NumPy (.npy). Pour un passage à l'échelle (>100k films), "
     "une vector database (FAISS, Qdrant, Milvus) serait nécessaire. "
     "Avec plus de temps : migration vers FAISS pour ANN search."),
    ("Poids de scoring fixes",
     "Les poids (35/25/20/15/5) sont définis manuellement. Un tuning basé sur du feedback "
     "utilisateur ou une optimisation bayésienne serait plus rigoureux."),
    ("Auth simplifiée",
     "Un seul utilisateur démo (admin/admin). En production : base utilisateurs, "
     "hashing bcrypt individuel, refresh token rotation."),
]
for title, desc in weaknesses:
    p = doc.add_paragraph()
    bold_run(p, f"✗ {title} — ", color=RGBColor(0xC0, 0x39, 0x2B))
    p.add_run(desc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  12. LIMITES ET PISTES D'AMÉLIORATION
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Limites et Pistes d'Amélioration", level=1)

doc.add_heading("12.1 Améliorations IA prioritaires", level=2)
improvements = [
    "Vector Database (FAISS / Qdrant) — indexation ANN pour passage à l'échelle",
    "Cross-encoder re-ranking — bi-encoder (SBERT) pour recall, cross-encoder pour precision",
    "Fine-tuning SBERT — entraînement sur paires (query, film) annotées cinéma",
    "Évaluation formelle — Precision@K, Recall@K, MRR, NDCG sur set de test",
    "Calibration dynamique des poids — optimisation bayésienne ou A/B testing",
    "Détection de drift — monitoring de la distribution des scores dans le temps",
]
for imp in improvements:
    p = doc.add_paragraph()
    bold_run(p, "→ ", color=BLUE)
    p.add_run(imp)

doc.add_heading("12.2 Améliorations produit/engineering", level=2)
for imp in [
    "MLOps pipeline (DVC pour versionnage data/modèle, CI/CD pour tests)",
    "Tests automatisés API (pytest + fixtures) et tests d'intégration",
    "Dashboard monitoring (latence, cache ratio, erreurs provider, usage)",
    "Support multilingue (SBERT multilingue + détection de langue)",
    "Chatbot conversationnel pour affiner les préférences itérativement",
]:
    p = doc.add_paragraph()
    bold_run(p, "→ ", color=BLUE)
    p.add_run(imp)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  13. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("13. Conclusion", level=1)

add_para(
    "CineMatch démontre la construction d'un pipeline IA sémantique complet, "
    "de la donnée brute CSV à la recommandation expliquée par IA Générative. "
    "Le projet couvre l'intégralité de la chaîne de valeur data/IA :",
    bold_phrases=["pipeline IA sémantique complet"],
)

skills = [
    "NLP appliqué — embeddings SBERT, similarité cosinus, prompt engineering structuré",
    "Ingénierie ML — pipeline offline/online, scoring multicritère, normalisation min-max",
    "GenAI contrôlée — RAG grounding, cache intelligent, multi-provider, coût maîtrisé",
    "Software Engineering — API REST typée, auth JWT, OpenAPI, modularité Python",
    "Vision produit — UX explicative, radar charts, presets, historique, page architecture",
]
for skill in skills:
    p = doc.add_paragraph()
    bold_run(p, "• ", color=BLUE)
    p.add_run(skill)

doc.add_paragraph()
add_para(
    "Ce projet illustre les compétences visées par le RNCP Bloc 2 : "
    "pilotage et implémentation d'une solution IA exploitable, argumentée et évolutive. "
    "Nous avons appris à articuler théorie (embeddings, cosinus, RAG) et pratique "
    "(API, cache, UI, démo), tout en maintenant une gouvernance responsable.",
    bold_phrases=["RNCP Bloc 2", "gouvernance responsable"],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  14. ANNEXES
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("14. Annexes", level=1)

doc.add_heading("Annexe A — Arborescence projet", level=2)
add_code("""EFREI-Project_AI-M1/
├── src/
│   ├── api/           # FastAPI app, routers, schemas
│   ├── preprocessing/ # block_builder, referential_builder
│   ├── embedding/     # sbert_loader, embedding_builder
│   ├── scoring/       # coverage_scorer (weighted formula)
│   ├── recommendation/# recommender (Top-3 engine)
│   ├── genAi/         # llm_client, prompt_builder, preset_queries
│   ├── user_profile/  # profile_encoder, questionnaire_schema
│   ├── core/          # config, security, llm_runtime_config, history
│   └── services/      # movie_catalog
├── frontend/          # Next.js 16 + React + Tailwind
├── data/              # raw CSV + processed CSV + SQLite history
├── models/embeddings/ # 4 × NPY matrices (9841 × 384)
└── docs/              # OpenAPI specs, rapport""")

doc.add_heading("Annexe B — Prompts GenAI utilisés", level=2)
doc.add_paragraph("1. Prompt d'enrichissement (build_enrichment_prompt) — voir §7.3")
doc.add_paragraph("2. Prompt de justification (build_cinewatch_prompt) — voir §7.3")

doc.add_heading("Annexe C — Variables d'environnement", level=2)
add_table(
    ["Variable", "Rôle", "Valeur par défaut"],
    [
        ["SECRET_KEY", "Clé signature JWT", "dev-key"],
        ["LLM_PROVIDER", "Provider actif", "ollama"],
        ["LLM_URL", "Endpoint Ollama", "http://localhost:11434/api/generate"],
        ["LLM_MODEL", "Modèle Ollama", "phi3:mini"],
        ["ANTHROPIC_API_KEY", "Clé Claude", "(vide)"],
        ["GEMINI_API_KEY", "Clé Gemini", "(vide)"],
        ["LLM_CACHE_TTL", "TTL cache (secondes)", "86400 (24h)"],
        ["LLM_CACHE_MAX_SIZE", "Max entrées cache", "200"],
    ],
)

doc.add_heading("Annexe D — Dépendances Python", level=2)
add_table(
    ["Package", "Rôle"],
    [
        ["sentence-transformers", "Embeddings SBERT"],
        ["torch", "Backend PyTorch pour SBERT"],
        ["scikit-learn", "Cosine similarity"],
        ["pandas / numpy", "Manipulation données + calcul matriciel"],
        ["fastapi / uvicorn", "API REST + serveur ASGI"],
        ["python-jose[cryptography]", "JWT create/decode"],
        ["bcrypt", "Hashing mot de passe"],
        ["requests", "Appels HTTP (Ollama, Claude, Gemini)"],
        ["matplotlib", "Radar charts backend (Streamlit)"],
    ],
)

# ═══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════════════════════════
doc.save(str(OUT))
print(f"Report saved to: {OUT}")
print(f"Size: {OUT.stat().st_size / 1024:.0f} KB")
