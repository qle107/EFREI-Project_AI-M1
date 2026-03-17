"""
Generate the final CineMatch project report in English as a professional .docx Word document.
Blue-themed, AI-focused, with Gantt chart, GitHub workflow, critique section.
"""
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT = Path(__file__).resolve().parent.parent / "docs" / "FINAL_REPORT_AI_GENERATIVE_EN.docx"

BLUE = RGBColor(0x1B, 0x4F, 0x72)
BLUE_LIGHT = RGBColor(0x21, 0x6F, 0xA3)
BLUE_ACCENT = RGBColor(0x2E, 0x86, 0xC1)
DARK = RGBColor(0x1C, 0x1C, 0x2E)
GRAY = RGBColor(0x5D, 0x6D, 0x7E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Global styles ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = DARK
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for lvl, sz in [(1, Pt(22)), (2, Pt(16)), (3, Pt(13))]:
    hs = doc.styles[f"Heading {lvl}"]
    hs.font.name = "Calibri"
    hs.font.size = sz
    hs.font.color.rgb = BLUE
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(18 if lvl == 1 else 12)
    hs.paragraph_format.space_after = Pt(6)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def bold_run(para, text, *, color=None):
    r = para.add_run(text)
    r.font.bold = True
    if color:
        r.font.color.rgb = color
    return r


def add_para(text, *, bold_phrases=None, after_pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after_pt)
    if bold_phrases:
        remaining = text
        for bp in bold_phrases:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                p.add_run(remaining[:idx])
            bold_run(p, bp, color=BLUE)
            remaining = remaining[idx + len(bp):]
        if remaining:
            p.add_run(remaining)
    else:
        p.add_run(text)
    return p


def add_code(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1C, 0x1C, 0x2E)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB"/>')
    p._element.get_or_add_pPr().append(shading)


def add_img_placeholder(caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"[ INSERT IMAGE: {caption} ]")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = BLUE_ACCENT
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB"/>')
    p._element.get_or_add_pPr().append(shading)


def set_cell_bg(cell, hex_color: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def add_table(headers, rows, *, header_bg="1B4F72", header_fg=WHITE):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = header_fg
        r.font.name = "Calibri"
        set_cell_bg(cell, header_bg)
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg = "FFFFFF" if ri % 2 == 0 else "EBF5FB"
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = "Calibri"
            set_cell_bg(cell, bg)
    doc.add_paragraph()


def blue_separator():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 60)
    run.font.color.rgb = BLUE_ACCENT
    run.font.size = Pt(8)


# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

add_img_placeholder("EFREI Paris logo — file: Efrei-logo-couleur.svg")
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CineMatch")
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Semantic and Generative Movie Recommendation Engine")
r.font.size = Pt(16)
r.font.color.rgb = BLUE_LIGHT

doc.add_paragraph()
blue_separator()
doc.add_paragraph()

cover_lines = [
    ("Master Data Engineering and AI — EFREI Paris", False),
    ("Academic year 2025–2026", False),
    ("", False),
    ("CHUEMBOU MBAH Adrien Duval  &  QUAND DAT Le", True),
    ("", False),
    ("Supervisor: Sarah Malaeb", False),
    ("", False),
    ("Certifying project RNCP40875", True),
    ("Expert in Data Engineering — Block 2:", True),
    ("Steering and implementation of AI solutions", True),
]
for text, is_bold in cover_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    if is_bold:
        r.font.bold = True
        r.font.color.rgb = BLUE

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "2. Executive Summary",
    "3. Introduction and Context",
    "4. User Need Analysis",
    "5. Methodology and Project Management",
    "6. Data Reference",
    "7. AI Pipeline and Architecture",
    "8. Technical Implementation",
    "9. User Interface / Prototype",
    "10. Results and Testing",
    "11. Project Critique — Strengths and Weaknesses",
    "12. Limitations and Future Improvements",
    "13. Conclusion",
    "14. Appendices",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.color.rgb = BLUE_LIGHT

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  2. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Executive Summary", level=1)

add_para(
    "CineMatch is a movie recommendation engine combining semantic NLP "
    "(Sentence-BERT, cosine similarity, weighted multi-criteria scoring) and Generative AI "
    "(LLM via Ollama / Claude / Gemini). The system turns a user query in natural language "
    "into Top-3 recommendations with a structured explanation and a generated cinephile profile.",
    bold_phrases=["semantic NLP", "Generative AI", "Top-3", "cinephile profile"],
)

add_para(
    "The architecture relies on two phases: an offline pipeline (building the semantic "
    "reference + SBERT encoding of 9,841 films on 4 independent axes) and an online pipeline "
    "(user profiling, cosine similarity, weighted aggregation, ranking, GenAI enrichment "
    "and explanation). The whole is exposed via a secured FastAPI (JWT) and a Next.js frontend.",
    bold_phrases=["offline pipeline", "online pipeline", "9,841 films", "4 independent axes"],
)

doc.add_heading("Key technologies", level=3)
add_table(
    ["Component", "Technology", "Role"],
    [
        ["Embeddings", "all-MiniLM-L6-v2 (SBERT)", "Dense 384-dim semantic encoding"],
        ["Similarity", "scikit-learn cosine_similarity", "User ↔ corpus comparison per axis"],
        ["Scoring", "CineMatch weighted formula", "0.35 mood + 0.25 theme + 0.20 style + 0.15 desc + 0.05 recency"],
        ["GenAI", "Ollama / Claude / Gemini", "Short text enrichment + explanation + cinephile profile"],
        ["LLM cache", "SHA-256 in-memory, TTL 24h", "Cost/latency reduction (req 1.5.3)"],
        ["API", "FastAPI + JWT HS-256", "REST, OpenAPI, authentication"],
        ["Frontend", "Next.js 16 + React + Tailwind", "Interactive prototype, SVG radar charts"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. INTRODUCTION AND CONTEXT
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Introduction and Context", level=1)

doc.add_heading("3.1 Problem Statement", level=2)
add_para(
    "Classical recommendation systems (collaborative filtering, content-based on metadata) "
    "struggle to capture the deep semantics of user preferences. A viewer asking for "
    "\"a dark psychological thriller with moral dilemmas\" cannot be satisfied by a simple "
    "Genre = Thriller filter. The system must understand the meaning of the text, not just keywords.",
    bold_phrases=["deep semantics", "understand the meaning of the text"],
)

add_para(
    "CineMatch addresses this by projecting films and queries into a shared semantic vector "
    "space via Sentence-BERT, then measuring proximity on 4 independent axes: Mood (atmosphere), "
    "Theme (genre/theme), Style (narrative pacing), and Description (free text). This multi-axis "
    "design enables fine-grained, explainable scoring via radar charts.",
    bold_phrases=["shared semantic vector space", "Sentence-BERT", "4 independent axes",
                  "Mood", "Theme", "Style", "Description"],
)

doc.add_heading("3.2 Why This AI Approach", level=2)
add_para(
    "Choosing a semantic-embedding approach over classical filtering is justified by several "
    "fundamental reasons:",
)

rationale = [
    ("Dense vs. sparse representation",
     "SBERT embeddings produce dense 384-dimensional vectors that capture semantic relations "
     "between concepts (\"dark\" ≈ \"noir\" ≈ \"gritty\"), whereas bag-of-words or TF-IDF would see no link."),
    ("Universal comparability",
     "By encoding users and films in the same vector space, cosine similarity becomes a natural "
     "relevance metric, with no need for ad-hoc rules."),
    ("Explainability by axis",
     "Splitting into 4 independent semantic dimensions (rather than a single combined vector) "
     "allows explaining why a film is recommended: \"high mood score but medium style.\" This is the basis for the radar chart."),
    ("Controlled GenAI (RAG-like)",
     "The generative layer does not invent: it receives the mathematically computed Top-3 and "
     "per-axis scores, then produces a grounded explanation. This limits hallucinations and cost."),
]
for title, desc in rationale:
    p = doc.add_paragraph()
    bold_run(p, f"{title} — ", color=BLUE)
    p.add_run(desc)

doc.add_heading("3.3 Alternatives Considered", level=2)
add_table(
    ["Approach", "Advantages", "Why not chosen"],
    [
        ["Collaborative filtering", "No need for film features", "Requires massive user history (cold start)"],
        ["TF-IDF + cosine", "Simple, fast", "Sparse representation, no semantic capture"],
        ["all-mpnet-base-v2", "Better STS score (86.8)", "4× slower, 768 dims, overkill for short movie text"],
        ["Fine-tuned SBERT", "Maximum quality", "Requires annotated (query/film) dataset, out of scope"],
        ["LLM-only (no retrieval)", "Simplicity", "Expensive, uncontrolled hallucinations, no scoring"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. USER NEED ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("4. User Need Analysis", level=1)

doc.add_heading("4.1 Target Persona", level=2)
add_para(
    "The typical user is an occasional to confirmed movie enthusiast, looking for personalized "
    "recommendations without a specific title in mind. They express their preferences via a "
    "combination of free text (\"I want a dark film with twists\") and structured choices "
    "(mood, genre, style, intensity via Likert 1–5 sliders).",
    bold_phrases=["free text", "structured choices", "Likert 1–5 sliders"],
)

doc.add_heading("4.2 Use Scenarios", level=2)
add_table(
    ["Scenario", "Input", "AI behaviour"],
    [
        ["A — Detailed", "\"Dark SF thriller with AI moral dilemmas, tense 5/5\"",
         "Direct SBERT encoding → cosine → Top-3 → explanation"],
        ["B — Short (< 5 words)", "\"sad movie\"",
         "GenAI enrichment → new encoding → cosine → Top-3"],
        ["C — Demo preset", "Click on 'Dark Crime Thriller'",
         "Full pipeline, LLM cache after 1st call → instant response"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. METHODOLOGY AND PROJECT MANAGEMENT
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Methodology and Project Management", level=1)

doc.add_heading("5.1 Adopted Approach", level=2)
add_para(
    "We followed a hybrid Kanban/Agile approach, with short functional increments and "
    "regular validation. Each batch delivers a testable artefact (pipeline, API, UI, GenAI, "
    "history, settings, documentation).",
    bold_phrases=["hybrid Kanban/Agile", "functional increments"],
)

doc.add_heading("5.2 Gantt Chart — Project Schedule", level=2)
add_para(
    "The table below summarizes the schedule by work package over ~8 weeks:"
)
add_table(
    ["Package", "Task", "Wk 1–2", "Wk 3–4", "Wk 5–6", "Wk 7–8"],
    [
        ["1", "Data cleaning + reference", "████", "", "", ""],
        ["2", "SBERT embedding pipeline", "████", "", "", ""],
        ["3", "Weighted multi-criteria scoring", "", "████", "", ""],
        ["4", "FastAPI REST API + JWT", "", "████", "", ""],
        ["5", "GenAI layer (prompts, cache, multi-provider)", "", "", "████", ""],
        ["6", "Next.js frontend (questionnaire, results, radar)", "", "", "████", ""],
        ["7", "Presets, history, runtime settings", "", "", "", "████"],
        ["8", "Documentation, report, presentation", "", "", "", "████"],
    ],
)

add_img_placeholder("Visual Gantt chart (if graphic version available)")

doc.add_heading("5.3 Pair Collaboration with GitHub", level=2)
add_para(
    "Work was organized on a shared GitHub repository. Collaboration followed these practices:",
    bold_phrases=["shared GitHub repository"],
)

collab_points = [
    ("Git versioning",
     "All source files (backend, frontend, scripts, data processing) are versioned. "
     "Each major feature (pipeline, API, GenAI, settings) was committed separately."),
    ("Modular organization",
     "Code is structured in separate Python modules (src/preprocessing, src/embedding, "
     "src/scoring, src/genAi, src/api) to avoid merge conflicts."),
    ("Task split",
     "The pair divided work by domain: one member focused on the AI pipeline "
     "(embeddings, scoring, GenAI), the other on API and frontend."),
    ("Ongoing documentation",
     "PROJECT_DOCUMENTATION.md was kept up to date throughout the project, "
     "serving as a shared reference between both members."),
    (".env and configuration",
     "Sensitive variables (API keys, JWT secrets) are managed via .env with a "
     "versioned .env.example for easier onboarding."),
]
for title, desc in collab_points:
    p = doc.add_paragraph()
    bold_run(p, f"• {title} — ", color=BLUE)
    p.add_run(desc)

doc.add_heading("5.4 Risk Management", level=2)
add_table(
    ["Identified risk", "Impact", "Mitigation"],
    [
        ["LLM latency (Ollama on CPU)", "High", "SHA-256 cache (TTL 24h) + switch to Claude/Gemini cloud"],
        ["Cloud API unavailability", "Medium", "Automatic fallback to local Ollama"],
        ["Semantic noise (\"neutral\" labels)", "Medium", "4-axis decomposition + differentiated weighting"],
        ["Embedding quality drift", "Low", "Log monitoring, optional offline re-build"],
        ["Git conflicts in pair", "Low", "Separate modules, regular communication"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. DATA REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Data Reference", level=1)

doc.add_heading("6.1 Source and Volume", level=2)
add_para(
    "The corpus is a TMDB-derived dataset with 9,841 films and fields: "
    "Title, Overview (synopsis), Genre, Release_Date, Vote_Average, Poster_Url. "
    "The dataset was cleaned (duplicate removal, year extraction, genre normalization) "
    "to produce Database_Cleaned.csv.",
    bold_phrases=["9,841 films", "Database_Cleaned.csv"],
)

doc.add_heading("6.2 Semantic Block Extraction", level=2)
add_para(
    "The block_builder.py module turns each film overview into 4 semantic blocks "
    "via a keyword taxonomy with regex (word-boundary) matching. The process is "
    "deterministic and reproducible.",
    bold_phrases=["4 semantic blocks", "keyword taxonomy"],
)

add_table(
    ["Block", "Source", "Taxonomy", "Example labels"],
    [
        ["Mood", "Overview", "MOOD_KEYWORDS", "dark, uplifting, tense, heroic"],
        ["Theme", "Overview", "THEME_KEYWORDS", "crime, fantasy, survival, family"],
        ["NarrativeStyle", "Overview", "STYLE_KEYWORDS", "action, mystery, drama"],
        ["Description", "Overview + Genre", "—", "Full text enriched with genre"],
    ],
)

doc.add_heading("How extraction works", level=3)
add_para(
    "For each film, the overview is lowercased and scanned with regex "
    "\\b{keyword}\\b (word boundary). If a keyword from a category is detected, the "
    "corresponding label is added. A film can have several labels per block (e.g. \"dark, tense\"). "
    "If no keyword matches, the default label is \"neutral\".",
    bold_phrases=["regex", "word boundary", "\"neutral\""],
)

add_code("""def extract_block(text, taxonomy):
    text = clean_text(text)  # lowercase + handle NaN
    detected = []
    for label, keywords in taxonomy.items():
        for word in keywords:
            if re.search(rf"\\b{word}\\b", text):
                detected.append(label)
                break  # one match per label is enough
    return ", ".join(detected) if detected else "neutral"

# Applied to 9,841 films:
df["Mood"]  = df["Overview"].apply(lambda x: extract_block(x, MOOD_KEYWORDS))
df["Theme"] = df["Overview"].apply(lambda x: extract_block(x, THEME_KEYWORDS))
df["NarrativeStyle"] = df["Overview"].apply(
    lambda x: extract_block(x, STYLE_KEYWORDS)
)
df["Description"] = df["Overview"] + " Genre: " + df["Genre"]""")

doc.add_heading("6.3 Produced Reference Schema", level=2)
add_table(
    ["Column", "Type", "Description"],
    [
        ["FilmID", "int", "Unique identifier (0-based index)"],
        ["Title", "str", "Film title"],
        ["Mood", "str", "Mood labels (e.g. \"dark, tense\")"],
        ["Theme", "str", "Theme labels (e.g. \"crime, family\")"],
        ["NarrativeStyle", "str", "Style labels (e.g. \"mystery, action\")"],
        ["Description", "str", "Overview + \" Genre: \" + Genre"],
        ["release_year", "int", "Release year"],
    ],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. AI PIPELINE AND ARCHITECTURE (MAJOR SECTION)
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("7. AI Pipeline and Architecture", level=1)

add_para(
    "The technical core of CineMatch is a 7-step pipeline, split into two phases: "
    "an offline phase (run once to build the index) and an online phase "
    "(run on every user request). Each step is detailed below with "
    "corresponding code, technical choices, and rationale.",
    bold_phrases=["7-step pipeline", "offline phase", "online phase"],
)

add_img_placeholder(
    "Full pipeline diagram: Questionnaire → Preprocessing → "
    "SBERT Encoding → Cosine Similarity → Weighted Scoring → Top-3 → GenAI Explanation"
)

# ── 7.1 OFFLINE ──
doc.add_heading("7.1 Offline Phase — Building the Semantic Index", level=2)

doc.add_heading("Step 1: Semantic block extraction", level=3)
add_para(
    "See section 6.2. referential_builder.py reads Database_Cleaned.csv, applies "
    "block_builder.build_semantic_blocks(), and saves movies_referential.csv.",
)

doc.add_heading("Step 2: SBERT Encoding of the Corpus (Embedding Builder)", level=3)
add_para(
    "Each block of each film is encoded independently with the SBERT model "
    "all-MiniLM-L6-v2 to produce a dense 384-dimensional vector. "
    "The result is 4 NumPy matrices of shape (9,841 × 384), one per semantic axis.",
    bold_phrases=["all-MiniLM-L6-v2", "384 dimensions", "4 NumPy matrices", "(9,841 × 384)"],
)

add_para(
    "Why encode each axis separately? A single embedding per film "
    "(e.g. concatenating title+overview+genre) would lose the granularity needed for "
    "per-axis scoring and radar visualization. By separating axes, we can "
    "weight mood (35%) differently from description (15%) and explain "
    "precisely where the film matches or diverges.",
    bold_phrases=["Why encode each axis separately?", "weight differently"],
)

add_code("""# src/embedding/embedding_builder.py
from sentence_transformers import SentenceTransformer
import numpy as np

model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

# Encode each semantic axis independently
mood_emb  = model.encode(df["Mood"].tolist(),  convert_to_numpy=True)
theme_emb = model.encode(df["Theme"].tolist(), convert_to_numpy=True)
style_emb = model.encode(df["NarrativeStyle"].tolist(), convert_to_numpy=True)
desc_emb  = model.encode(df["Description"].tolist(),    convert_to_numpy=True)

# Save as NumPy arrays for fast loading at inference
np.save("models/embeddings/mood_embeddings.npy",  mood_emb)   # (9841, 384)
np.save("models/embeddings/theme_embeddings.npy", theme_emb)  # (9841, 384)
np.save("models/embeddings/style_embeddings.npy", style_emb)  # (9841, 384)
np.save("models/embeddings/desc_embeddings.npy",  desc_emb)   # (9841, 384)""")

add_para(
    "The SBERT model works as follows: text is tokenized, passed through a MiniLM "
    "backbone (6-layer transformer), then mean-pooled to produce a fixed-size vector. "
    "Semantically similar texts are projected to nearby regions in 384-dim space.",
    bold_phrases=["MiniLM", "mean-pooled", "fixed-size vector"],
)

# ── 7.2 ONLINE ──
doc.add_heading("7.2 Online Phase — Real-Time Recommendation", level=2)

doc.add_heading("Step 3: User Profile Encoding", level=3)
add_para(
    "The user's hybrid questionnaire (free text + structured choices + Likert) "
    "is translated into 4 text blocks via structured templates. Each block is then "
    "encoded with the same SBERT model as the corpus, ensuring comparability in "
    "the same vector space.",
    bold_phrases=["4 text blocks", "same SBERT model", "same vector space"],
)

add_code("""# src/user_profile/profile_encoder.py
def build_profile_blocks(self, questionnaire):
    mood_block = (
        f"The user wants a {questionnaire.preferred_mood} mood "
        f"with intensity level {questionnaire.mood_intensity}"
    )
    theme_block = (
        f"The user prefers {questionnaire.preferred_genre} themes "
        f"with interest level {questionnaire.theme_interest}"
    )
    style_block = (
        f"The user enjoys {questionnaire.preferred_style} narrative "
        f"pacing with interest level {questionnaire.style_interest}"
    )
    desc_block = questionnaire.description
    if questionnaire.preferred_era:
        desc_block += f" Preferred era: {questionnaire.preferred_era}."
    if questionnaire.preferred_director:
        desc_block += f" Style of: {questionnaire.preferred_director}."
    return {"mood": mood_block, "theme": theme_block,
            "style": style_block, "description": desc_block}

# Encode with same SBERT model
profile = {k: model.encode(v) for k, v in blocks.items()}""")

add_para(
    "Likert intensity (e.g. \"intensity level 5\" vs \"intensity level 1\") affects "
    "the position of the vector in semantic space. Although subtle, this variation "
    "changes similarity scores in a way consistent with user intent.",
    bold_phrases=["Likert intensity", "affects the vector position"],
)

doc.add_heading("Step 4: Cosine Similarity Computation", level=3)
add_para(
    "For each axis, cosine similarity is computed between the user vector "
    "and the 9,841 corresponding film vectors. The formula is:",
    bold_phrases=["cosine similarity"],
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("sim(A, B) = (A · B) / (‖A‖ × ‖B‖)")
r.font.bold = True
r.font.size = Pt(12)
r.font.color.rgb = BLUE

add_para(
    "The result lies in [−1, 1] (typically [0, 1] for normalized embeddings). "
    "The higher the value, the more semantically close the film is to the profile on that axis.",
    bold_phrases=["[−1, 1]", "semantically close"],
)

add_code("""# src/scoring/coverage_scorer.py
from sklearn.metrics.pairwise import cosine_similarity

# 4 independent comparisons: user vector vs corpus matrix
mood_sim  = cosine_similarity([user_profile["mood"]],  self.mood_emb)[0]
theme_sim = cosine_similarity([user_profile["theme"]], self.theme_emb)[0]
style_sim = cosine_similarity([user_profile["style"]], self.style_emb)[0]
desc_sim  = cosine_similarity([user_profile["description"]], self.desc_emb)[0]
# Each result: array of 9,841 scores""")

doc.add_heading("Step 5: Min-Max Normalization per Axis", level=3)
add_para(
    "Before aggregation, each axis is normalized independently to [0, 1] "
    "via min-max scaling. This ensures axes with naturally different distributions "
    "contribute fairly to the final score.",
    bold_phrases=["normalized independently", "[0, 1]", "min-max scaling"],
)

add_code("""def normalize(self, sim):
    return (sim - sim.min()) / (sim.max() - sim.min() + 1e-8)
    # epsilon 1e-8 to avoid division by zero""")

doc.add_heading("Step 6: Weighted Aggregation — CineMatch Formula", level=3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
r = p.add_run(
    "CoverageScore = 0.35 × Mood + 0.25 × Theme + 0.20 × Style + 0.15 × Desc + 0.05 × Recency"
)
r.font.bold = True
r.font.size = Pt(12)
r.font.color.rgb = BLUE
shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB"/>')
p._element.get_or_add_pPr().append(shading)

add_code("""# Weighted aggregation
final_score = (
    0.35 * mood_sim +     # Atmosphere is most discriminating
    0.25 * theme_sim +    # Genre/theme is second filter
    0.20 * style_sim +    # Narrative pacing refines matching
    0.15 * desc_sim +     # Free text captures nuances
    0.05 * recency_score  # Light tiebreaker for freshness
)""")

doc.add_heading("Detailed Weight Justification", level=3)
add_table(
    ["Axis", "Weight", "Why"],
    [
        ["Mood", "35%", "Emotional atmosphere is the most discriminating criterion: a \"dark\" vs \"uplifting\" film changes the experience radically."],
        ["Theme", "25%", "Genre is the second natural filter: a sci-fi fan does not want romance."],
        ["Style", "20%", "Narrative pacing refines: a \"slow-burn\" vs \"fast-paced\" thriller is not interchangeable."],
        ["Description", "15%", "Free text captures nuances not covered by structured axes (director, era, specific tone)."],
        ["Recency", "5%", "Light tiebreaker: exp(−Δyear/10). A 10-year-old film scores ~0.37, 30 years ~0.05. Never dominant."],
    ],
)

add_para(
    "Recency score is computed with exponential decay:",
    bold_phrases=["exponential decay"],
)
add_code("""recency_score = np.exp(-(current_year - release_year) / 10)""")

doc.add_heading("Step 7: Top-3 Ranking and Output", level=3)
add_code("""ranked = self.df.sort_values(by="CoverageScore", ascending=False)
top3 = ranked.head(3)
# Each movie includes: FilmID, Title, CoverageScore,
# MoodScore, ThemeScore, StyleScore, DescScore""")

add_para(
    "The 3 films with the highest CoverageScore are returned with their detailed "
    "sub-scores. These sub-scores feed the SVG radar chart in the frontend and let "
    "the user see on which axis the film matches best.",
    bold_phrases=["detailed sub-scores", "SVG radar chart"],
)

doc.add_page_break()

# ── 7.3 GENAI ──
doc.add_heading("7.3 GenAI Layer — Enrichment and Explanation", level=2)

add_para(
    "The generative layer intervenes in a controlled way at two precise points in the pipeline. "
    "It never replaces the mathematical scoring: it complements it in natural language.",
    bold_phrases=["controlled", "never replaces the mathematical scoring"],
)

doc.add_heading("Short-Text Enrichment (< 5 words)", level=3)
add_para(
    "When the user enters a very short description (e.g. \"sad movie\"), "
    "a single LLM call expands it into 2–3 sentences before SBERT encoding. "
    "This improves the description vector quality without changing other axes.",
    bold_phrases=["single LLM call", "before SBERT encoding"],
)

add_code("""# src/genAi/prompt_builder.py
def build_enrichment_prompt(short_description: str) -> str:
    return (
        "You are a movie preference assistant. The user typed a very "
        "brief description. Expand this into a detailed 2-3 sentence "
        "description that preserves the original intent and adds "
        "useful context about tone, themes, pacing, and style.\\n\\n"
        f'User input: "{short_description}"\\n\\n'
        "Expanded description (2-3 sentences, no preamble):"
    )""")

doc.add_heading("Structured Explanation + Cinephile Profile", level=3)
add_para(
    "After Top-3 ranking, a single LLM call receives user preferences "
    "and the 3 films with their per-axis scores. The prompt requests a structured "
    "output in sections: OVERVIEW, FILM_1, FILM_2, FILM_3, REFINE, CINEPHILE_PROFILE.",
    bold_phrases=["single LLM call", "structured output"],
)

add_code("""# Prompt structure (condensed)
prompt = (
    "You are a cinematic AI analyst. Write a structured explanation "
    "using EXACTLY these headers:\\n"
    "OVERVIEW: [why these 3 match]\\n"
    "FILM_1: {title1} [strongest dimensions]\\n"
    "FILM_2: {title2} [...]\\n"
    "FILM_3: {title3} [...]\\n"
    "REFINE: [how to get better results]\\n"
    "CINEPHILE_PROFILE: [2-sentence taste profile]\\n\\n"
    "USER PREFERENCES: ...\\n"
    "SEMANTIC SCORES: ...\\n"
    "RULES: Do NOT invent data. Stay grounded in scores."
)""")

doc.add_heading("LLM Cache — Cost and Latency Reduction", level=3)
add_para(
    "Each prompt is hashed via SHA-256 and the response stored in memory with TTL (24h) "
    "and max size (200 entries). When an identical prompt is resubmitted, the response is "
    "served instantly from cache. This mechanism is essential for presets "
    "and repeated demos.",
    bold_phrases=["SHA-256", "TTL (24h)", "200 entries", "served instantly"],
)

add_code("""# src/genAi/llm_client.py — PromptCache (simplified)
class _PromptCache:
    def __init__(self, ttl=86400, max_size=200):
        self._store = {}  # {sha256_key: (response, timestamp)}

    def get(self, prompt):
        key = sha256(prompt)
        if key in store and not expired:
            return cached_response  # CACHE HIT

    def put(self, prompt, response):
        if at_capacity:
            evict_oldest()
        store[key] = (response, now())""")

doc.add_heading("Multi-Provider LLM", level=3)
add_table(
    ["Provider", "Model", "Typical latency", "Cost", "Use"],
    [
        ["Ollama (local)", "phi3:mini", "2–10s (CPU)", "Free", "Development, no internet"],
        ["Claude (Anthropic)", "claude-3-5-haiku", "1–3s", "Paid API", "Higher quality"],
        ["Gemini (Google)", "gemini-1.5-flash", "1–2s", "Paid API", "Fast, good quality"],
    ],
)

add_para(
    "The active provider can be switched at runtime from the Settings page or the "
    "\"Explanations by\" selector on the Recommend page, without restarting the server.",
    bold_phrases=["switched at runtime", "without restarting"],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. TECHNICAL IMPLEMENTATION
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Technical Implementation", level=1)

doc.add_heading("8.1 SBERT Model Comparison", level=2)
add_table(
    ["Criterion", "all-MiniLM-L6-v2 ✓", "all-mpnet-base-v2", "multi-qa-MiniLM"],
    [
        ["Parameters", "22M", "109M", "22M"],
        ["Vector dimension", "384", "768", "384"],
        ["STS Benchmark", "82.6", "86.8", "81.0"],
        ["CPU speed", "~4000 sent/s", "~1000 sent/s", "~4000 sent/s"],
        ["Best use", "Short/medium text", "High STS quality", "Q&A"],
        ["Fit for project?", "Yes (50–200 word overviews)", "Too slow for 9841 films", "No (Q&A oriented)"],
    ],
)

doc.add_heading("8.2 Hybrid Questionnaire → AI Vectors", level=2)
add_table(
    ["User field", "Type", "AI transformation"],
    [
        ["description", "Free text (min 10 chars)", "Direct SBERT encoding → Description axis"],
        ["preferred_mood", "Chip selection", "Template \"The user wants a {mood} mood...\" → Mood axis"],
        ["preferred_genre", "Chip selection", "Template \"The user prefers {genre} themes...\" → Theme axis"],
        ["preferred_style", "Chip selection", "Template \"The user enjoys {style} pacing...\" → Style axis"],
        ["mood_intensity", "Likert 1–5", "Included in mood template → influences vector"],
        ["theme_interest", "Likert 1–5", "Included in theme template → influences vector"],
        ["style_interest", "Likert 1–5", "Included in style template → influences vector"],
        ["preferred_era", "Optional", "Appended to desc_block → influences Description axis"],
        ["preferred_director", "Optional", "Appended to desc_block → influences Description axis"],
    ],
)

doc.add_heading("8.3 Main API Endpoints", level=2)
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["POST", "/api/v1/recommendations", "JWT", "Top-3 + GenAI explanation"],
        ["GET", "/api/v1/recommendations/presets", "—", "List available presets"],
        ["POST", "/api/v1/recommendations/presets/{id}", "JWT", "Run a preset"],
        ["GET", "/api/v1/recommendations/history", "JWT", "History per user"],
        ["GET/PUT", "/api/v1/settings/llm", "JWT", "Read/update LLM config"],
        ["POST", "/api/v1/settings/llm/clear-cache", "JWT", "Purge LLM cache"],
    ],
)

doc.add_heading("8.4 Governance and Responsible Use", level=2)
gov_points = [
    "LLM cache with TTL and max size — controls cost and latency",
    "JWT authentication on all sensitive endpoints",
    "Full logging of generations in SQLite — traceability",
    "Structured logs for LLM calls (provider, latency, cache hit/miss)",
    "Runtime provider switch without restart — resilience",
    "Clear separation: data / AI logic / API service / interface",
    "API keys never returned by the API (stored in memory only)",
]
for point in gov_points:
    p = doc.add_paragraph()
    bold_run(p, "• ", color=BLUE)
    p.add_run(point)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. USER INTERFACE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("9. User Interface — Prototype", level=1)

add_para(
    "The CineMatch prototype is built with Next.js 16, React, and Tailwind CSS. "
    "The interface guides the user through a 5-step questionnaire, displays "
    "Top-3 results with radar charts, GenAI explanation, and cinephile profile.",
    bold_phrases=["Next.js 16", "5 steps", "radar charts"],
)

screens = [
    "CineMatch home page — stats (9841 films, 384-dim, 4 axes) + login",
    "Recommendation questionnaire — free description + mood/genre/style chips + Likert sliders",
    "LLM provider selection (Ollama/Claude/Gemini) at top of Recommend page",
    "Top-3 results — film posters, CoverageScore, structured explanation, cinephile profile",
    "SVG radar charts per film — Mood, Theme, Style, Description",
    "How It Works page — full pipeline architecture, code excerpts, scoring formula",
    "Settings page — provider, temperature, model, cache TTL/size, cache purge",
    "Recommendation history — list of past generations with timestamp",
]
for s in screens:
    add_img_placeholder(s)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  10. RESULTS AND TESTING
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Results and Testing", level=1)

doc.add_heading("10.1 Demonstrations by Profile", level=2)
add_table(
    ["Test profile", "Mood", "Genre", "Intensity", "Expected Top-1 type", "Enriched?", "Cached?"],
    [
        ["Dark Crime Thriller", "dark", "crime", "5/5", "Noir/thriller", "No", "Yes (2nd call)"],
        ["Feel-Good Family", "uplifting", "family", "3/5", "Family/adventure", "No", "Yes (2nd call)"],
        ["Mind-Bending Sci-Fi", "tense", "sci-fi", "4/5", "Cerebral SF", "No", "No"],
        ["\"sad movie\" (short)", "—", "—", "—", "Emotional drama", "Yes", "No"],
    ],
)

doc.add_heading("10.2 API Response Metadata", level=2)
indicators = [
    ("coverage_score / mood_score / theme_score / style_score / desc_score",
     "Numeric scores [0, 1] per axis — feed the radar chart"),
    ("explanation", "Structured GenAI text (OVERVIEW, FILM_1/2/3, REFINE) grounded in scores"),
    ("cinephile_profile", "2 sentences describing the user's taste"),
    ("description_enriched", "true if short text was enriched by GenAI before encoding"),
    ("cached", "true if explanation came from cache (no new LLM call)"),
    ("llm_provider", "Provider used: ollama / anthropic / gemini"),
]
for name, desc in indicators:
    p = doc.add_paragraph()
    bold_run(p, f"{name} — ", color=BLUE)
    p.add_run(desc)

add_img_placeholder("Radar charts compared for 3 different profiles (dark vs family vs sci-fi)")
add_img_placeholder("Example JSON response with all fields (coverage_score, explanation, etc.)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  11. PROJECT CRITIQUE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Project Critique — Strengths and Weaknesses", level=1)

doc.add_heading("11.1 Strengths", level=2)
strengths = [
    ("End-to-end AI pipeline",
     "From raw CSV data to explained recommendation via GenAI, the project covers "
     "the full chain: preprocessing → embeddings → scoring → ranking → "
     "explanation. This is a major asset to demonstrate mastery of the ML pipeline."),
    ("Explainable multi-axis scoring",
     "Splitting into 4 independent axes (Mood, Theme, Style, Description) makes the score "
     "interpretable. The radar chart lets the user see where the match is strong "
     "and where it is weak. Few recommendation systems offer this level of transparency."),
    ("Grounded and controlled GenAI",
     "Generative AI never runs unchecked: it receives the computed scores and "
     "mathematically selected films. The prompt includes explicit RULES "
     "(\"Do NOT invent data\"). This is responsible RAG implementation."),
    ("Well-designed LLM cache",
     "The SHA-256 + TTL + eviction mechanism addresses real cost and latency issues. "
     "Presets benefit especially from this cache for smooth demos."),
    ("Switchable multi-provider",
     "Being able to switch between Ollama (local/free), Claude, and Gemini in real time, "
     "without restart, is a significant operational advantage."),
    ("Polished user interface",
     "The 5-step questionnaire, SVG radar charts, cinephile profile, "
     "How It Works page, and Settings form a complete prototype."),
]
for title, desc in strengths:
    p = doc.add_paragraph()
    bold_run(p, f"✓ {title} — ", color=BLUE)
    p.add_run(desc)

doc.add_heading("11.2 Weaknesses and Improvement Areas", level=2)
weaknesses = [
    ("Limited lexical taxonomy",
     "Block extraction relies on keyword regex. An overview with no known keyword "
     "gets the \"neutral\" label, reducing discrimination. "
     "With more time: NER or fine-tuned classification could replace this heuristic."),
    ("No SBERT fine-tuning",
     "The model is used out-of-the-box. Fine-tuning on (user query, relevant film) pairs "
     "would significantly improve quality. "
     "With more resources: build an annotated dataset + training."),
    ("Mostly qualitative evaluation",
     "We did not implement supervised metrics (Precision@K, NDCG, MRR). "
     "Evaluation relies on manual testing and result inspection. "
     "With more time: formal benchmark with annotators."),
    ("Static, English-only dataset",
     "The 9,841 films are fixed and in English. No real-time ingestion or multilingual support. "
     "With more resources: live TMDB API + multilingual SBERT model."),
    ("No dedicated vector database",
     "Embeddings are stored as NumPy (.npy). For scale (>100k films), "
     "a vector database (FAISS, Qdrant, Milvus) would be needed. "
     "With more time: migration to FAISS for ANN search."),
    ("Fixed scoring weights",
     "Weights (35/25/20/15/5) are set manually. Tuning based on user feedback "
     "or Bayesian optimization would be more rigorous."),
    ("Simplified auth",
     "Single demo user (admin/admin). In production: user base, "
     "per-user bcrypt hashing, refresh token rotation."),
]
for title, desc in weaknesses:
    p = doc.add_paragraph()
    bold_run(p, f"✗ {title} — ", color=RGBColor(0xC0, 0x39, 0x2B))
    p.add_run(desc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  12. LIMITATIONS AND FUTURE IMPROVEMENTS
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Limitations and Future Improvements", level=1)

doc.add_heading("12.1 Priority AI Improvements", level=2)
improvements = [
    "Vector database (FAISS / Qdrant) — ANN indexing for scale",
    "Cross-encoder re-ranking — bi-encoder (SBERT) for recall, cross-encoder for precision",
    "SBERT fine-tuning — training on annotated (query, film) movie pairs",
    "Formal evaluation — Precision@K, Recall@K, MRR, NDCG on test set",
    "Dynamic weight calibration — Bayesian optimization or A/B testing",
    "Drift detection — monitoring score distribution over time",
]
for imp in improvements:
    p = doc.add_paragraph()
    bold_run(p, "→ ", color=BLUE)
    p.add_run(imp)

doc.add_heading("12.2 Product/Engineering Improvements", level=2)
for imp in [
    "MLOps pipeline (DVC for data/model versioning, CI/CD for tests)",
    "Automated API tests (pytest + fixtures) and integration tests",
    "Monitoring dashboard (latency, cache ratio, provider errors, usage)",
    "Multilingual support (multilingual SBERT + language detection)",
    "Conversational chatbot to refine preferences iteratively",
]:
    p = doc.add_paragraph()
    bold_run(p, "→ ", color=BLUE)
    p.add_run(imp)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  13. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("13. Conclusion", level=1)

add_para(
    "CineMatch demonstrates the construction of a complete semantic AI pipeline, "
    "from raw CSV data to recommendation explained by Generative AI. "
    "The project covers the full data/AI value chain:",
    bold_phrases=["complete semantic AI pipeline"],
)

skills = [
    "Applied NLP — SBERT embeddings, cosine similarity, structured prompt engineering",
    "ML engineering — offline/online pipeline, multi-criteria scoring, min-max normalization",
    "Controlled GenAI — RAG grounding, smart cache, multi-provider, managed cost",
    "Software engineering — typed REST API, JWT auth, OpenAPI, Python modularity",
    "Product vision — explanatory UX, radar charts, presets, history, architecture page",
]
for skill in skills:
    p = doc.add_paragraph()
    bold_run(p, "• ", color=BLUE)
    p.add_run(skill)

doc.add_paragraph()
add_para(
    "This project illustrates the skills targeted by RNCP Block 2: "
    "steering and implementation of an exploitable, well-argued, and evolvable AI solution. "
    "We learned to combine theory (embeddings, cosine, RAG) and practice "
    "(API, cache, UI, demo), while maintaining responsible governance.",
    bold_phrases=["RNCP Block 2", "responsible governance"],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  14. APPENDICES
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_heading("14. Appendices", level=1)

doc.add_heading("Appendix A — Project Structure", level=2)
add_code("""EFREI-Project_AI-M1/
├── src/
│   ├── api/           # FastAPI app, routers, schemas
│   ├── preprocessing/ # block_builder, referential_builder
│   ├── embedding/     # sbert_loader, embedding_builder
│   ├── scoring/       # coverage_scorer (weighted formula)
│   ├── recommendation/# recommender (Top-3 engine)
│   ├── genAi/         # llm_client, prompt_builder, preset_queries
│   ├── user_profile/  # profile_encoder, questionnaire_schema
│   ├── core/          # config, security, llm_runtime_config, history
│   └── services/      # movie_catalog
├── frontend/          # Next.js 16 + React + Tailwind
├── data/              # raw CSV + processed CSV + SQLite history
├── models/embeddings/ # 4 × NPY matrices (9841 × 384)
└── docs/              # OpenAPI specs, report""")

doc.add_heading("Appendix B — GenAI Prompts Used", level=2)
doc.add_paragraph("1. Enrichment prompt (build_enrichment_prompt) — see §7.3")
doc.add_paragraph("2. Justification prompt (build_cinematch_prompt) — see §7.3")

doc.add_heading("Appendix C — Environment Variables", level=2)
add_table(
    ["Variable", "Role", "Default value"],
    [
        ["SECRET_KEY", "JWT signing key", "dev-key"],
        ["LLM_PROVIDER", "Active provider", "ollama"],
        ["LLM_URL", "Ollama endpoint", "http://localhost:11434/api/generate"],
        ["LLM_MODEL", "Ollama model", "phi3:mini"],
        ["ANTHROPIC_API_KEY", "Claude key", "(empty)"],
        ["GEMINI_API_KEY", "Gemini key", "(empty)"],
        ["LLM_CACHE_TTL", "Cache TTL (seconds)", "86400 (24h)"],
        ["LLM_CACHE_MAX_SIZE", "Max cache entries", "200"],
    ],
)

doc.add_heading("Appendix D — Python Dependencies", level=2)
add_table(
    ["Package", "Role"],
    [
        ["sentence-transformers", "SBERT embeddings"],
        ["torch", "PyTorch backend for SBERT"],
        ["scikit-learn", "Cosine similarity"],
        ["pandas / numpy", "Data handling + matrix computation"],
        ["fastapi / uvicorn", "REST API + ASGI server"],
        ["python-jose[cryptography]", "JWT create/decode"],
        ["bcrypt", "Password hashing"],
        ["requests", "HTTP calls (Ollama, Claude, Gemini)"],
        ["matplotlib", "Radar charts backend (Streamlit)"],
    ],
)

# ═══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════════════════════════
doc.save(str(OUT))
print(f"Report saved to: {OUT}")
print(f"Size: {OUT.stat().st_size / 1024:.0f} KB")
